from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from rest_framework import serializers
from ..models import FaqCategory, FaqArticle, FaqFile


class FaqCategorySerializer(serializers.ModelSerializer):
    articles_count = serializers.SerializerMethodField()

    class Meta:
        model = FaqCategory
        fields = ['id', 'name', 'order', 'articles_count', 'created_at']

    def get_articles_count(self, obj):
        return obj.articles.count()


class FaqArticleSerializer(serializers.ModelSerializer):
    author_name = serializers.SerializerMethodField()
    can_delete  = serializers.SerializerMethodField()

    class Meta:
        model = FaqArticle
        fields = ['id', 'category', 'title', 'content', 'author', 'author_name',
                  'order', 'created_at', 'updated_at', 'can_delete']
        read_only_fields = ['author', 'created_at', 'updated_at']

    def get_author_name(self, obj):
        if obj.author:
            return obj.author.full_name or obj.author.email
        return '—'

    def get_can_delete(self, obj):
        request = self.context.get('request')
        if not request:
            return False
        user = request.user
        is_admin = user.is_superuser or (hasattr(user, 'role') and user.role and user.role.name == 'admin')
        return is_admin or obj.author_id == user.id


class FaqCategoryViewSet(viewsets.ModelViewSet):
    queryset = FaqCategory.objects.all()
    serializer_class = FaqCategorySerializer
    permission_classes = [IsAuthenticated]

    def destroy(self, request, *args, **kwargs):
        user = request.user
        is_admin = user.is_superuser or (hasattr(user, 'role') and user.role and user.role.name == 'admin')
        if not is_admin:
            return Response({'error': 'Только администратор может удалять категории'}, status=403)
        return super().destroy(request, *args, **kwargs)


class FaqArticleViewSet(viewsets.ModelViewSet):
    serializer_class = FaqArticleSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = FaqArticle.objects.select_related('author', 'category')
        category_id = self.request.query_params.get('category')
        if category_id:
            qs = qs.filter(category_id=category_id)
        search = self.request.query_params.get('search')
        if search:
            qs = qs.filter(title__icontains=search) | qs.filter(content__icontains=search)
            qs = qs.distinct()
        return qs

    def perform_create(self, serializer):
        serializer.save(author=self.request.user)

    def destroy(self, request, *args, **kwargs):
        article = self.get_object()
        user = request.user
        is_admin = user.is_superuser or (hasattr(user, 'role') and user.role and user.role.name == 'admin')
        if not is_admin and article.author_id != user.id:
            return Response({'error': 'Можно удалять только свои статьи'}, status=403)
        return super().destroy(request, *args, **kwargs)

    def get_serializer_context(self):
        ctx = super().get_serializer_context()
        ctx['request'] = self.request
        return ctx


class FaqFileSerializer(serializers.ModelSerializer):
    url = serializers.SerializerMethodField()

    class Meta:
        model = FaqFile
        fields = ['id', 'name', 'size', 'url', 'created_at']

    def get_url(self, obj):
        request = self.context.get('request')
        if request:
            return request.build_absolute_uri(obj.file.url)
        return obj.file.url


def _is_admin(user):
    return user.is_superuser or (hasattr(user, 'role') and user.role and user.role.name == 'admin')

def _can_edit_article(user, article):
    """Редактировать статью может автор или администратор."""
    return _is_admin(user) or article.author_id == user.id

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 МБ

def _check_real_size(file, max_bytes=MAX_FILE_SIZE):
    """Считает реальный размер файла по содержимому, не по Content-Length."""
    size = 0
    for chunk in file.chunks():
        size += len(chunk)
        if size > max_bytes:
            return None  # превышен лимит
    file.seek(0)
    return size


class FaqFileView(APIView):
    """GET /api/clients/faq-articles/{id}/files/ — список файлов
       POST — загрузить файл
       DELETE /api/clients/faq-files/{id}/ — удалить файл
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, article_id):
        try:
            article = FaqArticle.objects.get(pk=article_id)
        except FaqArticle.DoesNotExist:
            return Response({'error': 'Статья не найдена'}, status=404)
        files = FaqFile.objects.filter(article=article)
        return Response(FaqFileSerializer(files, many=True, context={'request': request}).data)

    def post(self, request, article_id):
        try:
            article = FaqArticle.objects.get(pk=article_id)
        except FaqArticle.DoesNotExist:
            return Response({'error': 'Статья не найдена'}, status=404)
        if not _can_edit_article(request.user, article):
            return Response({'error': 'Нет прав для загрузки файлов в эту статью'}, status=403)
        file = request.FILES.get('file')
        if not file:
            return Response({'error': 'Файл не передан'}, status=400)
        real_size = _check_real_size(file)
        if real_size is None:
            return Response({'error': 'Файл слишком большой (максимум 10 МБ)'}, status=400)
        faq_file = FaqFile.objects.create(
            article=article,
            file=file,
            name=_safe_filename(file.name),
            size=real_size,
            uploaded_by=request.user,
        )
        return Response(FaqFileSerializer(faq_file, context={'request': request}).data, status=201)


class FaqFileDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, file_id):
        try:
            faq_file = FaqFile.objects.get(pk=file_id)
        except FaqFile.DoesNotExist:
            return Response({'error': 'Файл не найден'}, status=404)
        if not _is_admin(request.user) and faq_file.uploaded_by_id != request.user.id:
            return Response({'error': 'Нет прав для удаления'}, status=403)
        faq_file.file.delete(save=False)
        faq_file.delete()
        return Response({'ok': True})


def _safe_filename(filename):
    """Убирает path traversal и null-bytes, оставляет только безопасное имя файла."""
    import os, re
    # Убираем null-bytes и path separators
    filename = filename.replace('\x00', '').replace('\r', '').replace('\n', '')
    # Берём только базовое имя без пути
    filename = os.path.basename(filename)
    # Оставляем только безопасные символы
    filename = re.sub(r'[^\w\s\-_\.]', '', filename).strip()
    return filename or 'file'


def faq_image_path(filename):
    import uuid, os
    safe = _safe_filename(filename)
    ext = os.path.splitext(safe)[1].lower()
    # Разрешаем только известные расширения изображений
    if ext not in {'.jpg', '.jpeg', '.png', '.gif', '.webp'}:
        ext = '.bin'
    return f'faq/images/{uuid.uuid4().hex}{ext}'


class FaqImageUploadView(APIView):
    """POST /api/clients/faq-articles/{id}/images/ — загрузить картинку, вернуть URL"""
    permission_classes = [IsAuthenticated]

    def post(self, request, article_id):
        try:
            article = FaqArticle.objects.get(pk=article_id)
        except FaqArticle.DoesNotExist:
            return Response({'error': 'Статья не найдена'}, status=404)
        if not _can_edit_article(request.user, article):
            return Response({'error': 'Нет прав для загрузки изображений в эту статью'}, status=403)
        file = request.FILES.get('image')
        if not file:
            return Response({'error': 'Файл не передан'}, status=400)
        real_size = _check_real_size(file)
        if real_size is None:
            return Response({'error': 'Файл слишком большой (максимум 10 МБ)'}, status=400)

        # Проверяем по magic bytes — реальное содержимое, не заголовок
        header = file.read(16)
        file.seek(0)

        MAGIC = {
            b'\xff\xd8\xff':    ('jpg', 'image/jpeg'),
            b'\x89PNG\r\n':     ('png', 'image/png'),
            b'GIF87a':          ('gif', 'image/gif'),
            b'GIF89a':          ('gif', 'image/gif'),
            b'RIFF':            ('webp', 'image/webp'),  # webp начинается с RIFF
        }
        detected = None
        for magic, (ext, mime) in MAGIC.items():
            if header.startswith(magic):
                detected = (ext, mime)
                break
        # WEBP дополнительная проверка (RIFF....WEBP)
        if header[:4] == b'RIFF' and header[8:12] != b'WEBP':
            detected = None

        if not detected:
            return Response({'error': 'Разрешены только изображения (jpg, png, gif, webp)'}, status=400)

        real_ext, _ = detected
        import os
        from django.core.files.storage import default_storage
        import uuid as _uuid
        fname = f'faq/images/{_uuid.uuid4().hex}.{real_ext}'
        saved = default_storage.save(fname, file)
        url = request.build_absolute_uri(f'/media/{saved}')
        return Response({'url': url}, status=201)


class FaqImportView(APIView):
    """POST /api/clients/faq-articles/{id}/import/ — импорт из docx или pdf в HTML"""
    permission_classes = [IsAuthenticated]

    def post(self, request, article_id):
        try:
            article = FaqArticle.objects.get(pk=article_id)
        except FaqArticle.DoesNotExist:
            return Response({'error': 'Статья не найдена'}, status=404)
        if not _can_edit_article(request.user, article):
            return Response({'error': 'Нет прав для импорта в эту статью'}, status=403)
        file = request.FILES.get('file')
        if not file:
            return Response({'error': 'Файл не передан'}, status=400)

        name = file.name.lower()
        if name.endswith('.docx'):
            html = self._import_docx(file, article_id, request)
        elif name.endswith('.pdf'):
            html = self._import_pdf(file, article_id, request)
        else:
            return Response({'error': 'Поддерживаются только .docx и .pdf'}, status=400)

        return Response({'html': html})

    def _import_docx(self, file, article_id, request):
        from docx import Document
        from docx.oxml.ns import qn
        import uuid
        from django.core.files.base import ContentFile
        from django.core.files.storage import default_storage

        VML_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        doc = Document(file)

        # Собираем словарь rId → (bytes, ext)
        img_map = {}
        for rel in doc.part.rels.values():
            if 'image' in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    ct = rel.target_part.content_type
                    ext = ct.split('/')[-1].replace('jpeg', 'jpg')
                    img_map[rel.rId] = (img_data, ext)
                except Exception:
                    pass

        def save_image(img_data, ext):
            fname = f'faq/images/{uuid.uuid4().hex}.{ext}'
            default_storage.save(fname, ContentFile(img_data))
            return request.build_absolute_uri(f'/media/{fname}')

        def get_para_images(para_el):
            """Извлекает rId картинок из параграфа — оба формата: modern (a:blip) и VML (imagedata)"""
            urls = []
            # Современный формат
            for blip in para_el.findall('.//' + qn('a:blip')):
                rId = blip.get(qn('r:embed'))
                if rId and rId in img_map:
                    urls.append(save_image(*img_map[rId]))
            # VML формат (старый Word)
            for imgdata in para_el.findall('.//{urn:schemas-microsoft-com:vml}imagedata'):
                rId = imgdata.get(f'{{{VML_R}}}id')
                if rId and rId in img_map:
                    urls.append(save_image(*img_map[rId]))
            return urls

        html_parts = []
        import html as html_lib

        for para in doc.paragraphs:
            imgs = get_para_images(para._element)
            text = html_lib.escape(para.text.strip())
            style = para.style.name if para.style else ''

            for url in imgs:
                html_parts.append(f'<p><img src="{url}" style="max-width:100%;border-radius:6px;" /></p>')

            if text:
                if 'Heading 1' in style:
                    html_parts.append(f'<h2>{text}</h2>')
                elif 'Heading 2' in style or 'Heading 3' in style:
                    html_parts.append(f'<h3>{text}</h3>')
                else:
                    inner = ''
                    for run in para.runs:
                        t = html_lib.escape(run.text)
                        if not t:
                            continue
                        if run.bold and run.italic:
                            t = f'<strong><em>{t}</em></strong>'
                        elif run.bold:
                            t = f'<strong>{t}</strong>'
                        elif run.italic:
                            t = f'<em>{t}</em>'
                        elif run.underline:
                            t = f'<u>{t}</u>'
                        inner += t
                    if inner:
                        html_parts.append(f'<p>{inner}</p>')

        return ''.join(html_parts)

    def _import_pdf(self, file, article_id, request):
        import fitz
        import uuid
        from django.core.files.base import ContentFile
        from django.core.files.storage import default_storage

        pdf_bytes = file.read()
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
        html_parts = []
        global_seen_xrefs = set()
        import html as html_lib

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Собираем текстовые блоки с позицией
            items = []
            text_dict = page.get_text('dict')
            for block in text_dict['blocks']:
                if block['type'] == 0:
                    for line in block.get('lines', []):
                        text = ' '.join(s['text'] for s in line.get('spans', [])).strip()
                        if not text:
                            continue
                        sizes = [s['size'] for s in line.get('spans', []) if s.get('size')]
                        avg_size = sum(sizes) / len(sizes) if sizes else 12
                        y = block['bbox'][1]
                        items.append({'type': 'text', 'y': y, 'text': text, 'size': avg_size})

            # Собираем картинки с позицией
            # Дедупликация 1: глобальная по xref (одна картинка на весь документ)
            raw_imgs = []
            for img_info in page.get_image_info(xrefs=True):
                xref = img_info.get('xref', 0)
                if not xref or xref in global_seen_xrefs:
                    continue
                global_seen_xrefs.add(xref)
                bbox = img_info.get('bbox')
                y = bbox[1] if bbox else 0
                raw_imgs.append({'type': 'image', 'y': y, 'xref': xref})

            # Дедупликация 2: по близости bbox_y (оставляем с большей площадью)
            raw_imgs.sort(key=lambda x: x['y'])
            filtered_imgs = []
            for img in raw_imgs:
                # Проверяем не дубликат ли по Y-позиции (разница < 15px)
                is_dup = False
                for kept in filtered_imgs:
                    if abs(img['y'] - kept['y']) < 15:
                        # Оставляем с большей площадью
                        try:
                            d1 = doc.extract_image(img['xref'])
                            d2 = doc.extract_image(kept['xref'])
                            if d1.get('width',0)*d1.get('height',0) > d2.get('width',0)*d2.get('height',0):
                                filtered_imgs.remove(kept)
                                filtered_imgs.append(img)
                        except Exception:
                            pass
                        is_dup = True
                        break
                if not is_dup:
                    filtered_imgs.append(img)

            items.extend(filtered_imgs)

            # Сортируем по вертикальной позиции
            items.sort(key=lambda x: x['y'])

            for item in items:
                if item['type'] == 'text':
                    text = html_lib.escape(item['text'])
                    if item['size'] >= 16:
                        html_parts.append(f'<h2>{text}</h2>')
                    elif item['size'] >= 13:
                        html_parts.append(f'<h3>{text}</h3>')
                    else:
                        html_parts.append(f'<p>{text}</p>')
                elif item['type'] == 'image':
                    try:
                        img_data = doc.extract_image(item['xref'])
                        if not img_data:
                            continue
                        # Пропускаем слишком маленькие изображения (иконки, рамки)
                        w = img_data.get('width', 0)
                        h = img_data.get('height', 0)
                        if w < 30 or h < 30:
                            continue
                        ext = img_data.get('ext', 'png')
                        fname = f'faq/images/{uuid.uuid4().hex}.{ext}'
                        default_storage.save(fname, ContentFile(img_data['image']))
                        url = request.build_absolute_uri(f'/media/{fname}')
                        html_parts.append(f'<p><img src="{url}" style="max-width:100%;border-radius:6px;" /></p>')
                    except Exception:
                        pass

        doc.close()
        if not html_parts:
            return '<p>Не удалось извлечь содержимое из PDF.</p>'
        return ''.join(html_parts)


class FaqExportView(APIView):
    """GET /api/clients/faq-articles/{id}/export/?format=docx|pdf"""
    permission_classes = [IsAuthenticated]

    def get(self, request, article_id):
        try:
            article = FaqArticle.objects.select_related('author', 'category').get(pk=article_id)
        except FaqArticle.DoesNotExist:
            return Response({'error': 'Статья не найдена'}, status=404)

        fmt = request.query_params.get('type', 'docx').lower()
        if fmt == 'docx':
            return self._export_docx(article, request)
        elif fmt == 'pdf':
            return self._export_pdf(article, request)
        return Response({'error': 'Поддерживаются форматы: docx, pdf'}, status=400)

    def _export_docx(self, article, request):
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from html.parser import HTMLParser
        import io, os

        doc = Document()

        # Заголовок статьи
        title_para = doc.add_heading(article.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Метаданные
        meta = doc.add_paragraph()
        r1 = meta.add_run(f'Категория: {article.category.name}  |  Автор: {article.author.full_name if article.author else "—"}')
        r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r1.font.size = Pt(10)
        doc.add_paragraph()

        class HTMLToDocx(HTMLParser):
            def __init__(self):
                super().__init__()
                self.para = None
                self.bold = self.italic = self.underline = self.in_pre = False
                self.skip_depth = 0

            def handle_starttag(self, tag, attrs):
                attrs = dict(attrs)
                if tag in ('script', 'style'):
                    self.skip_depth += 1; return
                if self.skip_depth: return
                if tag in ('h1', 'h2'):
                    self.para = doc.add_heading('', level=1)
                elif tag == 'h3':
                    self.para = doc.add_heading('', level=2)
                elif tag == 'p':
                    self.para = doc.add_paragraph()
                elif tag == 'pre':
                    self.in_pre = True; self.para = doc.add_paragraph()
                elif tag == 'li':
                    self.para = doc.add_paragraph(style='List Bullet')
                elif tag in ('strong', 'b'): self.bold = True
                elif tag in ('em', 'i'): self.italic = True
                elif tag == 'u': self.underline = True
                elif tag == 'br':
                    if self.para: self.para.add_run('\n')
                elif tag == 'img':
                    src = attrs.get('src', '')
                    if not src:
                        return
                    img_para = doc.add_paragraph()
                    try:
                        if src.startswith('data:image/'):
                            import base64, io
                            from docx.shared import Inches, Pt, Emu
                            from PIL import Image as PILImage
                            header, b64data = src.split(',', 1)
                            img_bytes = base64.b64decode(b64data)
                            # Определяем реальный размер картинки
                            pil_img = PILImage.open(io.BytesIO(img_bytes))
                            orig_w, orig_h = pil_img.size  # в пикселях
                            dpi = pil_img.info.get('dpi', (96, 96))
                            dpi_x = dpi[0] if isinstance(dpi, tuple) else 96
                            # Переводим в дюймы
                            w_inches = orig_w / dpi_x
                            max_w = 5.5  # максимум 5.5 дюймов
                            if w_inches > max_w:
                                # Только уменьшаем — сохраняем пропорции
                                img_para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(max_w))
                            else:
                                # Оригинальный размер
                                img_para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(w_inches))
                        elif '/media/' in src:
                            from PIL import Image as PILImage
                            import io
                            rel = src.split('/media/')[-1].split('?')[0]
                            path = os.path.join('/app/media', rel)
                            if os.path.exists(path):
                                pil_img = PILImage.open(path)
                                orig_w, orig_h = pil_img.size
                                dpi = pil_img.info.get('dpi', (96, 96))
                                dpi_x = dpi[0] if isinstance(dpi, tuple) else 96
                                w_inches = orig_w / dpi_x
                                max_w = 5.5
                                if w_inches > max_w:
                                    img_para.add_run().add_picture(path, width=Inches(max_w))
                                else:
                                    img_para.add_run().add_picture(path, width=Inches(w_inches))
                            else:
                                img_para.add_run(f'[Изображение не найдено: {os.path.basename(src)}]')
                        self.para = None
                    except Exception:
                        img_para.add_run('[Изображение]')

            def handle_endtag(self, tag):
                if tag in ('script', 'style'):
                    self.skip_depth = max(0, self.skip_depth - 1); return
                if tag in ('strong', 'b'): self.bold = False
                elif tag in ('em', 'i'): self.italic = False
                elif tag == 'u': self.underline = False
                elif tag == 'pre': self.in_pre = False

            def handle_data(self, data):
                if self.skip_depth: return
                text = data
                if self.in_pre:
                    text = text.replace('Копировать', '').strip()
                    if not text: return
                    if self.para is None: self.para = doc.add_paragraph()
                    run = self.para.add_run(text)
                    run.font.name = 'Courier New'; run.font.size = Pt(10)
                    return
                if not text.strip(): return
                if self.para is None: self.para = doc.add_paragraph()
                run = self.para.add_run(text)
                run.bold = self.bold; run.italic = self.italic; run.underline = self.underline

        HTMLToDocx().feed(article.content)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)

        safe = ''.join(c for c in article.title if c.isalnum() or c in ' _-')[:50]
        from django.http import HttpResponse
        resp = HttpResponse(buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = f'attachment; filename="{safe}.docx"'
        return resp

    def _export_pdf(self, article, request):
        import html as hl
        from django.http import HttpResponse
        meta = f'{article.category.name} · {article.author.full_name if article.author else "—"}'
        css = """<style>
          @page{margin:1.5cm 2cm}
          body{font-family:Arial,sans-serif;font-size:13px;color:#333}
          h1{font-size:20px;margin-bottom:4px}
          h2{font-size:16px;margin-top:18px}
          h3{font-size:14px}
          .meta{color:#888;font-size:11px;margin-bottom:16px}
          img{max-width:100%}
          pre{background:#1e1e1e;color:#d4d4d4;padding:10px 14px;border-radius:6px;
              font-family:monospace;font-size:11px;white-space:pre-wrap;word-break:break-all}
          .copy-btn{display:none}
        </style>"""
        html_page = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>{hl.escape(article.title)}</title>{css}</head><body>
<h1>{hl.escape(article.title)}</h1>
<div class="meta">{hl.escape(meta)}</div><hr>
{article.content}
<script>window.onload=function(){{window.print()}}</script>
</body></html>"""
        return HttpResponse(html_page, content_type='text/html; charset=utf-8')
